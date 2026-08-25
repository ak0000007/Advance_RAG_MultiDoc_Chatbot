import uuid
from pathlib import Path

from docx import Document as DocxDocument

from src.models.document import Document
from src.ingestion.base import BaseLoader


class DOCXLoader(BaseLoader):

    def load(self, file_path: str):

        doc = DocxDocument(file_path)

        paragraphs = []

        for paragraph in doc.paragraphs:

            text = paragraph.text.strip()

            if text:
                paragraphs.append(text)

        content = "\n\n".join(paragraphs)

        path = Path(file_path)

        return Document(
            id=str(uuid.uuid4()),
            name=path.name,
            content=content,
            document_type="docx",
            source=str(path),
            metadata={
                "extension": path.suffix.lower(),
                "file_path": str(path)
            }
        )